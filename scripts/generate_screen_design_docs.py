"""Generate basic-design .docx files for every screen listed in SCREENS_LIST.md.

One file per screen, saved under docs/screen-designs/<module>/<slug>.docx.
Run:  python3 scripts/generate_screen_design_docs.py
"""

from __future__ import annotations

import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_ROOT = REPO_ROOT / "docs" / "screen-designs"


@dataclass
class Field:
    name: str
    type: str
    required: bool = False
    note: str = ""


@dataclass
class Screen:
    module: str
    module_label: str
    slug: str
    title_vn: str
    title_en: str
    route: str
    kind: str  # list | detail | create | edit | other
    purpose: str
    actors: list[str]
    roles: list[str]
    main_flow: list[str]
    alt_flow: list[str]
    layout: list[str]
    fields: list[Field]
    actions: list[str]
    business_rules: list[str]
    apis: list[str]
    validations: list[str]
    acceptance: list[str]
    open_questions: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Shared fragments
# ---------------------------------------------------------------------------

COMMON_ROLES = ["Admin", "Warehouse Manager", "Sales Staff", "Warehouse Staff", "Viewer"]

LIST_ACTIONS = [
    "Nút `+ Tạo mới` — điều hướng sang màn hình create.",
    "Nút `Import` — mở dialog upload Excel/CSV (nếu cho phép).",
    "Nút `Export` — xuất Excel toàn bộ hoặc theo filter hiện tại.",
    "Click dòng / nút `Chi tiết` — mở trang detail.",
    "Nút `Sửa` trên mỗi dòng — mở trang edit (nếu có).",
    "Nút `Xoá` trên mỗi dòng — confirm dialog, gọi API DELETE.",
    "Phân trang server-side (page, pageSize), mặc định 20 dòng/trang.",
]

DETAIL_ACTIONS = [
    "Nút `Sửa` — điều hướng sang trang edit (nếu có).",
    "Nút `Xoá` — confirm dialog, gọi API DELETE rồi back về list.",
    "Nút `Quay lại` — về list page.",
    "Nút `Xuất PDF` / `In` nếu có document (invoice / packing list).",
]

COMMON_ERROR_RULES = [
    "Áp dụng chuẩn `FRONTEND_ERROR_HANDLING_RULES.md`: toast error ngắn gọn, log console, giữ state form.",
    "Mọi request fail 401 ⇒ redirect `/login`.",
    "Mọi request fail 403 ⇒ toast `Bạn không có quyền` và disable action button.",
    "Validation client-side trước khi submit; hiển thị inline dưới field.",
]


def list_validation() -> list[str]:
    return [
        "Filter server-side, debounce 300ms.",
        "Sort theo cột cho phép; chỉ cột nằm trong allowlist mới gọi API.",
        "PageSize ∈ {10, 20, 50, 100}.",
    ]


def detail_validation() -> list[str]:
    return [
        "Load bằng `id` từ URL; nếu 404 ⇒ hiển thị trang `NotFound` và toast.",
        "Các trường hiển thị read-only, format theo locale VN (số, ngày, tiền tệ).",
        "Nếu dữ liệu liên kết bị xoá mềm ⇒ hiển thị badge `Đã xoá`.",
    ]


def create_validation() -> list[str]:
    return [
        "Bắt buộc điền đầy đủ trường `required`; disabled nút submit cho tới khi hợp lệ.",
        "Validate định dạng email, số, ngày, tiền tệ (> 0) bằng Zod/Yup.",
        "Chặn submit trùng (double-click) bằng state loading.",
        "Rollback data khi API lỗi; giữ nguyên input để user retry.",
    ]


def edit_validation() -> list[str]:
    return create_validation() + [
        "Khi load sẵn dữ liệu, so sánh với state hiện tại để bật/tắt nút `Lưu`.",
        "Cảnh báo rời trang nếu có thay đổi chưa lưu (beforeunload / route guard).",
    ]


# ---------------------------------------------------------------------------
# Screen definitions — one entry per screen in SCREENS_LIST.md
# ---------------------------------------------------------------------------


SCREENS: list[Screen] = []


def add(screen: Screen) -> None:
    SCREENS.append(screen)


# -- Dashboard ---------------------------------------------------------------

add(Screen(
    module="dashboard", module_label="Dashboard",
    slug="dashboard", title_vn="Trang chủ / Dashboard", title_en="Dashboard Home",
    route="/dashboard", kind="other",
    purpose="Cung cấp cái nhìn tổng quan về hoạt động kho, bán hàng, đơn hàng và cảnh báo tồn kho trong ngày.",
    actors=["Admin", "Warehouse Manager", "Sales Manager"],
    roles=["admin", "manager"],
    main_flow=[
        "User đăng nhập ⇒ được redirect vào `/dashboard`.",
        "Dashboard gọi song song các API KPI (sales today, inbound today, outbound today, low-stock count).",
        "Render các widget KPI card, biểu đồ sales 7/30 ngày, bảng cảnh báo tồn kho thấp.",
        "Click widget điều hướng về danh sách chi tiết tương ứng.",
    ],
    alt_flow=[
        "Nếu API KPI fail ⇒ widget hiển thị placeholder 'Không tải được dữ liệu', có nút Retry.",
        "Nếu user thuộc role 'Viewer' ⇒ ẩn các widget nhạy cảm (doanh thu).",
    ],
    layout=[
        "Header: filter khoảng thời gian (Today / 7 days / 30 days / Custom).",
        "Grid 4 KPI card hàng 1: doanh thu, đơn bán, phiếu nhập, phiếu xuất.",
        "Hàng 2: chart line sales, chart bar inbound/outbound.",
        "Hàng 3: bảng cảnh báo tồn kho thấp (top 10) + table notifications mới.",
    ],
    fields=[
        Field("kpiRevenueToday", "number", note="tổng doanh thu các quotation/po ký hôm nay"),
        Field("kpiOrdersToday", "number", note="số PO tạo hôm nay"),
        Field("kpiInboundToday", "number"),
        Field("kpiOutboundToday", "number"),
        Field("lowStockItems", "array<{product, qty, threshold}>"),
        Field("recentNotifications", "array<Notification>", note="5 mới nhất"),
    ],
    actions=[
        "Click card doanh thu ⇒ `/sales/quotations?date=today`.",
        "Click card phiếu nhập ⇒ `/warehouse/inbound?date=today`.",
        "Click dòng low-stock ⇒ `/master-data/products/[id]`.",
        "Click thông báo ⇒ `/notifications`.",
    ],
    business_rules=[
        "KPI tính theo timezone Asia/Ho_Chi_Minh.",
        "Low-stock = `stockQty <= product.reorderLevel`.",
    ],
    apis=[
        "GET /api/dashboard/kpis?range=today",
        "GET /api/dashboard/sales-trend?range=30d",
        "GET /api/dashboard/low-stock?limit=10",
        "GET /api/notifications?limit=5",
    ],
    validations=[
        "Range filter ∈ {today, 7d, 30d, custom}.",
        "Custom range: from <= to, khoảng tối đa 366 ngày.",
    ],
    acceptance=[
        "Load <= 1.5s với dữ liệu demo.",
        "Mỗi KPI hiển thị số và so sánh % so với kỳ trước.",
        "Role Viewer không thấy widget doanh thu.",
        "Responsive breakpoint >=1280px: 4 cột; >=960px: 2 cột; <960px: 1 cột.",
    ],
    open_questions=[
        "Có cần realtime (socket.io) cho low-stock không?",
        "KPI có cần export Excel?",
    ],
))


# -- Master Data: Suppliers --------------------------------------------------

add(Screen(
    module="master-data", module_label="Master Data",
    slug="suppliers-list", title_vn="Danh sách nhà cung cấp", title_en="Suppliers List",
    route="/master-data/suppliers", kind="list",
    purpose="Quản lý toàn bộ nhà cung cấp của doanh nghiệp: tìm kiếm, lọc, tạo, cập nhật, vô hiệu hoá.",
    actors=["Admin", "Purchasing Staff", "Warehouse Manager"],
    roles=["admin", "purchasing"],
    main_flow=[
        "User vào `/master-data/suppliers` ⇒ gọi `GET /api/suppliers`.",
        "Hiển thị bảng gồm cột: mã, tên, SĐT, email, quốc gia, trạng thái.",
        "User có thể filter theo trạng thái, quốc gia; search theo tên/mã.",
        "Click `+ Tạo mới` ⇒ mở Drawer tạo supplier.",
        "Click dòng ⇒ `/master-data/suppliers/[id]`.",
    ],
    alt_flow=[
        "Không có quyền create ⇒ ẩn nút `+ Tạo mới`.",
        "Import Excel fail validation ⇒ hiển thị bảng lỗi, cho tải file template.",
    ],
    layout=[
        "Header: search box + filter country + filter status + nút Import, Export, + Tạo mới.",
        "Body: DataGrid MUI (sortable, paginated).",
        "Footer: tổng số dòng, pagination.",
    ],
    fields=[
        Field("code", "string", required=True, note="duy nhất, uppercase"),
        Field("name", "string", required=True),
        Field("phone", "string"),
        Field("email", "email"),
        Field("country", "string"),
        Field("status", "enum(active, inactive)", note="filter"),
        Field("createdAt", "datetime"),
    ],
    actions=LIST_ACTIONS + ["Import Excel — validate + preview + confirm import."],
    business_rules=[
        "Không cho xoá supplier có purchase order đang active.",
        "Soft delete: cập nhật `status = inactive`, không xoá vật lý.",
    ],
    apis=[
        "GET /api/suppliers?q=&status=&country=&page=&pageSize=",
        "POST /api/suppliers",
        "POST /api/suppliers/import",
        "GET /api/suppliers/export.xlsx",
        "DELETE /api/suppliers/:id (soft)",
    ],
    validations=list_validation(),
    acceptance=[
        "Hiển thị đúng dữ liệu theo filter và pagination.",
        "Import Excel >= 1000 dòng chạy <5s; có progress.",
        "Nút + Tạo mới mở drawer, submit thành công thì reload list.",
        "Role không có quyền không thấy nút Create/Import/Delete.",
    ],
))


add(Screen(
    module="master-data", module_label="Master Data",
    slug="suppliers-detail", title_vn="Chi tiết nhà cung cấp", title_en="Supplier Detail",
    route="/master-data/suppliers/[id]", kind="detail",
    purpose="Hiển thị thông tin chi tiết nhà cung cấp, danh sách sản phẩm cung cấp, bảng giá và lịch sử giao dịch.",
    actors=["Admin", "Purchasing Staff"],
    roles=["admin", "purchasing"],
    main_flow=[
        "Từ list, click dòng supplier.",
        "Gọi `GET /api/suppliers/:id` (bao gồm products, pricing, recent POs).",
        "Hiển thị 3 tab: Thông tin, Bảng giá, Lịch sử PO.",
        "User có thể sửa inline hoặc mở drawer edit.",
    ],
    alt_flow=[
        "Supplier không tồn tại ⇒ trang NotFound.",
        "Supplier inactive ⇒ banner cảnh báo.",
    ],
    layout=[
        "Header: tên supplier, mã, status badge, nút Sửa / Xoá.",
        "Tab `Thông tin`: các field cơ bản, địa chỉ, liên hệ.",
        "Tab `Bảng giá`: table SupplierPricing.",
        "Tab `Lịch sử PO`: list 20 PO gần nhất.",
    ],
    fields=[
        Field("code", "string"), Field("name", "string"), Field("phone", "string"),
        Field("email", "email"), Field("country", "string"), Field("address", "text"),
        Field("taxCode", "string"), Field("contactPerson", "string"),
        Field("paymentTerm", "string"), Field("notes", "text"),
    ],
    actions=DETAIL_ACTIONS + ["Nút `Thêm sản phẩm cung cấp` ⇒ mở drawer."],
    business_rules=[
        "Chỉ Admin được xoá (soft).",
        "Lịch sử PO chỉ hiển thị 20 dòng mới nhất, có link xem tất cả.",
    ],
    apis=[
        "GET /api/suppliers/:id",
        "PATCH /api/suppliers/:id",
        "DELETE /api/suppliers/:id",
        "GET /api/suppliers/:id/pricing",
        "GET /api/suppliers/:id/purchase-orders?limit=20",
    ],
    validations=detail_validation(),
    acceptance=[
        "Mọi field hiển thị đúng format.",
        "Tab chuyển không gọi lại API đã cache.",
        "Nút Sửa/Xoá ẩn theo role.",
    ],
))


add(Screen(
    module="master-data", module_label="Master Data",
    slug="products-list", title_vn="Danh sách sản phẩm", title_en="Products List",
    route="/master-data/products", kind="list",
    purpose="Quản lý toàn bộ sản phẩm: SKU, mô tả, giá, tồn kho tổng, danh mục.",
    actors=["Admin", "Merchandiser", "Warehouse Manager"],
    roles=["admin", "merchandiser"],
    main_flow=[
        "User vào `/master-data/products` ⇒ `GET /api/products`.",
        "Hiển thị bảng: SKU, tên, danh mục, unit, giá bán, tồn kho.",
        "Filter theo category, status; search theo SKU/tên.",
        "Click dòng ⇒ chi tiết; `+ Tạo mới` ⇒ `/master-data/products/create`.",
    ],
    alt_flow=[
        "Không có sản phẩm ⇒ empty state với CTA `Tạo sản phẩm đầu tiên`.",
        "Import excel sai format ⇒ bảng lỗi row-by-row.",
    ],
    layout=[
        "Toolbar: search, filter category, filter status, Import, Export, + Tạo mới.",
        "Body: DataGrid với ảnh thumbnail + SKU + name + category + stockQty + price.",
        "Footer: pagination.",
    ],
    fields=[
        Field("sku", "string", required=True, note="duy nhất"),
        Field("name", "string", required=True),
        Field("categoryId", "uuid", required=True),
        Field("unit", "string", note="pcs/box/kg..."),
        Field("price", "decimal"),
        Field("stockQty", "number", note="tổng tồn từ inventory view"),
        Field("reorderLevel", "number"),
        Field("status", "enum(active,inactive)"),
    ],
    actions=LIST_ACTIONS + [
        "Import Excel sản phẩm.",
        "Nút `Xem tồn kho` ⇒ `/warehouse/inventory?productId=<id>`.",
    ],
    business_rules=[
        "SKU unique, viết hoa.",
        "Không cho xoá product đang có tồn kho > 0 hoặc có PO active.",
    ],
    apis=[
        "GET /api/products?q=&categoryId=&status=&page=",
        "POST /api/products",
        "POST /api/products/import",
        "GET /api/products/export.xlsx",
        "DELETE /api/products/:id",
    ],
    validations=list_validation(),
    acceptance=[
        "Lọc và search trả đúng kết quả.",
        "Thumbnail cache CDN; placeholder nếu thiếu.",
        "Import 5k rows < 10s có progress.",
    ],
))


add(Screen(
    module="master-data", module_label="Master Data",
    slug="products-detail", title_vn="Chi tiết sản phẩm", title_en="Product Detail",
    route="/master-data/products/[id]", kind="detail",
    purpose="Hiển thị toàn bộ thông tin sản phẩm kèm bảng giá nhà cung cấp và lịch sử xuất/nhập kho.",
    actors=["Admin", "Merchandiser"],
    roles=["admin", "merchandiser"],
    main_flow=[
        "Click dòng product từ list ⇒ `/master-data/products/[id]`.",
        "Gọi `GET /api/products/:id` (bao gồm pricing, stock, recent transactions).",
        "Hiển thị tab: Thông tin, Bảng giá NCC, Tồn kho theo kho, Lịch sử giao dịch.",
    ],
    alt_flow=[
        "Product không tồn tại ⇒ NotFound.",
        "Product inactive ⇒ hiển thị badge đỏ.",
    ],
    layout=[
        "Header: ảnh, SKU, tên, status, giá bán, nút Sửa/Xoá.",
        "Tab Thông tin: mô tả, kích thước, trọng lượng, barcode.",
        "Tab Bảng giá: table SupplierPricing theo sản phẩm.",
        "Tab Tồn kho: table theo warehouseId, availableQty.",
        "Tab Lịch sử: 50 giao dịch inbound/outbound gần nhất.",
    ],
    fields=[
        Field("sku", "string"), Field("name", "string"), Field("categoryId", "uuid"),
        Field("description", "richtext"), Field("unit", "string"), Field("price", "decimal"),
        Field("barcode", "string"), Field("dimensions", "string"), Field("weight", "number"),
        Field("reorderLevel", "number"), Field("safetyStock", "number"), Field("images", "array<url>"),
    ],
    actions=DETAIL_ACTIONS + [
        "Nút `Xem tồn kho` ⇒ inventory filter theo product.",
        "Nút `Thêm bảng giá NCC` ⇒ mở drawer.",
    ],
    business_rules=[
        "Giá bán phải >= chi phí gốc (cảnh báo nếu vi phạm).",
        "Không cho đổi SKU khi đã có transactions.",
    ],
    apis=[
        "GET /api/products/:id",
        "PATCH /api/products/:id",
        "DELETE /api/products/:id",
        "GET /api/products/:id/pricing",
        "GET /api/products/:id/inventory",
        "GET /api/products/:id/transactions?limit=50",
    ],
    validations=detail_validation(),
    acceptance=[
        "Mọi tab load đúng dữ liệu.",
        "Sửa và lưu thành công ⇒ refetch và refresh UI.",
        "Role không đủ quyền ⇒ disable Edit/Delete.",
    ],
))


add(Screen(
    module="master-data", module_label="Master Data",
    slug="product-categories-list", title_vn="Danh sách danh mục sản phẩm", title_en="Product Categories List",
    route="/master-data/product-categories", kind="list",
    purpose="Quản lý cây danh mục sản phẩm (có thể đa cấp).",
    actors=["Admin", "Merchandiser"],
    roles=["admin", "merchandiser"],
    main_flow=[
        "Load `/master-data/product-categories` ⇒ `GET /api/product-categories?tree=true`.",
        "Hiển thị TreeView 2 panel: cây bên trái, detail + danh sách product bên phải.",
        "User có thể thêm/sửa/xoá node, drag-drop để đổi parent.",
    ],
    alt_flow=[
        "Xoá node có products ⇒ chặn, hiển thị số lượng product đang dùng.",
    ],
    layout=[
        "Panel trái: tree danh mục với nút `+ thêm con` khi hover.",
        "Panel phải: detail của node được chọn + table product thuộc node.",
    ],
    fields=[
        Field("code", "string", required=True),
        Field("name", "string", required=True),
        Field("parentId", "uuid?"),
        Field("sortOrder", "number"),
        Field("status", "enum(active,inactive)"),
    ],
    actions=[
        "Thêm node con, đổi tên inline, xoá node.",
        "Drag-drop reorder / đổi parent.",
        "Nút `Export` xuất toàn bộ cây Excel.",
    ],
    business_rules=[
        "Tối đa 5 cấp.",
        "Không cho circular reference.",
    ],
    apis=[
        "GET /api/product-categories?tree=true",
        "POST /api/product-categories",
        "PATCH /api/product-categories/:id",
        "DELETE /api/product-categories/:id",
        "PATCH /api/product-categories/:id/move",
    ],
    validations=list_validation(),
    acceptance=[
        "Drag-drop cập nhật đúng parent và sortOrder.",
        "Xoá có cảnh báo khi node có children/products.",
    ],
))


add(Screen(
    module="master-data", module_label="Master Data",
    slug="product-categories-detail", title_vn="Chi tiết danh mục sản phẩm", title_en="Product Category Detail",
    route="/master-data/product-categories/[id]", kind="detail",
    purpose="Xem chi tiết danh mục, danh sách sản phẩm thuộc danh mục và danh mục con.",
    actors=["Admin", "Merchandiser"],
    roles=["admin", "merchandiser"],
    main_flow=[
        "Từ tree, click node ⇒ load `GET /api/product-categories/:id`.",
        "Hiển thị tab Thông tin, Danh mục con, Sản phẩm.",
    ],
    alt_flow=[
        "Không có sản phẩm ⇒ empty state với CTA thêm sản phẩm.",
    ],
    layout=[
        "Header: tên, code, breadcrumb cha.",
        "Tab Thông tin: các field, status.",
        "Tab Danh mục con: list + CTA thêm.",
        "Tab Sản phẩm: list product + filter.",
    ],
    fields=[
        Field("code", "string"), Field("name", "string"), Field("description", "text"),
        Field("parentId", "uuid?"), Field("sortOrder", "number"), Field("status", "enum(active,inactive)"),
    ],
    actions=DETAIL_ACTIONS + ["Nút `Thêm danh mục con`, `Thêm sản phẩm vào danh mục`."],
    business_rules=[
        "Không cho đổi parent thành chính nó hoặc descendants.",
    ],
    apis=[
        "GET /api/product-categories/:id",
        "PATCH /api/product-categories/:id",
        "DELETE /api/product-categories/:id",
        "GET /api/product-categories/:id/products",
    ],
    validations=detail_validation(),
    acceptance=[
        "Các tab load đúng dữ liệu.",
        "Sửa tên category ⇒ tree bên ngoài cập nhật realtime (socket/refetch).",
    ],
))


add(Screen(
    module="master-data", module_label="Master Data",
    slug="supplier-pricing-list", title_vn="Danh sách giá nhà cung cấp", title_en="Supplier Pricing List",
    route="/master-data/supplier-pricing", kind="list",
    purpose="Quản lý bảng giá theo nhà cung cấp × sản phẩm × khoảng thời gian.",
    actors=["Admin", "Purchasing Staff"],
    roles=["admin", "purchasing"],
    main_flow=[
        "Load `/master-data/supplier-pricing` ⇒ `GET /api/supplier-pricing`.",
        "Hiển thị table: supplier, product, giá, currency, đơn vị, hiệu lực từ-đến, trạng thái.",
        "User tạo/sửa/xoá record; có thể import Excel.",
    ],
    alt_flow=[
        "Giá có ngày hiệu lực chồng chéo ⇒ server trả 409, UI hiển thị modal giải thích.",
    ],
    layout=[
        "Toolbar: filter supplier, product, currency; search; + Tạo, Import, Export.",
        "Body: DataGrid sortable.",
    ],
    fields=[
        Field("supplierId", "uuid", required=True),
        Field("productId", "uuid", required=True),
        Field("price", "decimal", required=True),
        Field("currency", "enum(VND,USD,EUR...)", required=True),
        Field("unit", "string"),
        Field("moq", "number", note="min order qty"),
        Field("effectiveFrom", "date", required=True),
        Field("effectiveTo", "date"),
        Field("status", "enum(active,expired,upcoming)"),
    ],
    actions=LIST_ACTIONS,
    business_rules=[
        "Unique theo (supplierId, productId, currency, effectiveFrom).",
        "Không cho overlap khoảng hiệu lực cùng (supplier, product, currency).",
    ],
    apis=[
        "GET /api/supplier-pricing?supplierId=&productId=&currency=&page=",
        "POST /api/supplier-pricing",
        "PATCH /api/supplier-pricing/:id",
        "DELETE /api/supplier-pricing/:id",
        "POST /api/supplier-pricing/import",
    ],
    validations=list_validation() + [
        "effectiveFrom <= effectiveTo.",
        "price > 0.",
    ],
    acceptance=[
        "Tạo record trùng khoảng hiệu lực hiện báo lỗi rõ.",
        "Export Excel đủ cột; import Excel hiển thị preview.",
    ],
))


# -- Sales: Quotations -------------------------------------------------------

add(Screen(
    module="sales", module_label="Sales",
    slug="quotations-list", title_vn="Danh sách báo giá", title_en="Quotations List",
    route="/sales/quotations", kind="list",
    purpose="Quản lý các báo giá gửi cho khách hàng: tạo mới, theo dõi trạng thái duyệt, chuyển thành PO.",
    actors=["Sales Staff", "Sales Manager", "Admin"],
    roles=["sales", "sales-manager", "admin"],
    main_flow=[
        "Load `/sales/quotations` ⇒ `GET /api/quotations`.",
        "Filter theo status (draft, sent, approved, rejected, converted).",
        "Click `+ Tạo báo giá` ⇒ `/sales/quotations/create`.",
        "Click dòng ⇒ chi tiết.",
    ],
    alt_flow=[
        "Sales Staff chỉ thấy báo giá của mình; Manager thấy tất cả.",
    ],
    layout=[
        "Toolbar: search, filter status, filter customer, filter ngày; + Tạo, Export.",
        "Body: DataGrid với cột: mã, khách hàng, tổng tiền, status, ngày gửi, sales owner.",
    ],
    fields=[
        Field("code", "string"),
        Field("customerId", "uuid", required=True),
        Field("totalAmount", "decimal"),
        Field("currency", "enum"),
        Field("status", "enum(draft,sent,approved,rejected,converted)"),
        Field("sentAt", "datetime"),
        Field("ownerId", "uuid", note="sales owner"),
    ],
    actions=LIST_ACTIONS + [
        "Nút `Duplicate` clone thành báo giá draft mới.",
        "Nút `Gửi` set status=sent + gửi email.",
    ],
    business_rules=[
        "Chỉ owner hoặc manager được edit.",
        "Không cho xoá nếu đã `approved` hoặc `converted`.",
    ],
    apis=[
        "GET /api/quotations?status=&ownerId=&q=",
        "POST /api/quotations",
        "DELETE /api/quotations/:id",
        "POST /api/quotations/:id/send",
    ],
    validations=list_validation(),
    acceptance=[
        "Sales Staff không thấy báo giá của người khác.",
        "Filter status chuẩn xác.",
        "Duplicate tạo draft mới với items y hệt.",
    ],
))


add(Screen(
    module="sales", module_label="Sales",
    slug="quotations-create", title_vn="Tạo báo giá mới", title_en="Create Quotation",
    route="/sales/quotations/create", kind="create",
    purpose="Form tạo báo giá mới với khách hàng, danh sách sản phẩm, điều khoản, thuế, chiết khấu.",
    actors=["Sales Staff", "Sales Manager"],
    roles=["sales", "sales-manager"],
    main_flow=[
        "Sales chọn customer ⇒ autofill thông tin liên hệ, VAT.",
        "Thêm line items: product, qty, unit price, discount.",
        "Xem tổng tiền, thuế, tổng cuối.",
        "Lưu draft hoặc Gửi (status=sent).",
    ],
    alt_flow=[
        "Thiếu dữ liệu bắt buộc ⇒ disable nút Save/Send.",
        "Sản phẩm không đủ tồn kho ⇒ hiển thị cảnh báo vàng.",
    ],
    layout=[
        "Section 1: thông tin khách hàng + ngày hiệu lực.",
        "Section 2: bảng line items có nút Thêm sản phẩm, xoá dòng, tính subtotal.",
        "Section 3: terms (ghi chú, thanh toán, giao hàng).",
        "Section 4: Summary + nút Save draft, Gửi.",
    ],
    fields=[
        Field("customerId", "uuid", required=True),
        Field("validUntil", "date", required=True),
        Field("currency", "enum", required=True),
        Field("items", "array<{productId, qty, unitPrice, discount, note}>", required=True),
        Field("vatRate", "number", note="mặc định 10"),
        Field("shippingTerms", "text"),
        Field("paymentTerms", "text"),
        Field("notes", "richtext"),
    ],
    actions=[
        "Thêm sản phẩm từ search dialog.",
        "Lưu nháp.",
        "Gửi (chuyển sang status sent + gửi mail khách).",
        "Huỷ và quay lại list.",
    ],
    business_rules=[
        "Tổng tiền = Σ(qty × unitPrice × (1-discount)) × (1+vatRate).",
        "Chiết khấu dòng tối đa 50%; >30% cần manager duyệt.",
    ],
    apis=[
        "POST /api/quotations",
        "GET /api/customers?q=",
        "GET /api/products?q=&available=true",
    ],
    validations=create_validation() + [
        "Ít nhất 1 line item.",
        "qty > 0, unitPrice >= 0.",
    ],
    acceptance=[
        "Gửi báo giá thành công ⇒ quay về list và toast thành công.",
        "Tự tính tổng realtime.",
        "Validation inline hoạt động.",
    ],
))


add(Screen(
    module="sales", module_label="Sales",
    slug="quotations-detail", title_vn="Chi tiết báo giá", title_en="Quotation Detail",
    route="/sales/quotations/[id]", kind="detail",
    purpose="Xem thông tin báo giá, lịch sử gửi, trạng thái phê duyệt, in PDF hoặc chuyển thành PO.",
    actors=["Sales Staff", "Sales Manager", "Admin"],
    roles=["sales", "sales-manager", "admin"],
    main_flow=[
        "Click dòng quotation ⇒ detail page.",
        "Hiển thị thông tin, line items, lịch sử (audit trail).",
        "Nút `Gửi`, `Duyệt`, `Từ chối`, `Chuyển PO`, `In PDF`.",
    ],
    alt_flow=[
        "Báo giá đã converted ⇒ disable mọi hành động thay đổi.",
    ],
    layout=[
        "Header: mã, khách hàng, status badge, nút hành động.",
        "Section line items.",
        "Section terms.",
        "Section audit log.",
    ],
    fields=[
        Field("code", "string"), Field("customerId", "uuid"),
        Field("items", "array"), Field("totalAmount", "decimal"),
        Field("status", "enum"), Field("ownerId", "uuid"),
        Field("history", "array<{action, actor, at}>"),
    ],
    actions=DETAIL_ACTIONS + [
        "Nút `Gửi` (nếu draft).",
        "Nút `Duyệt`/`Từ chối` (manager).",
        "Nút `Chuyển PO` (nếu approved).",
        "Nút `In PDF` / `Gửi email`.",
    ],
    business_rules=[
        "Chỉ manager duyệt được báo giá đã gửi.",
        "Chuyển PO ⇒ tạo PO draft tham chiếu quotationId.",
    ],
    apis=[
        "GET /api/quotations/:id",
        "PATCH /api/quotations/:id",
        "POST /api/quotations/:id/approve",
        "POST /api/quotations/:id/reject",
        "POST /api/quotations/:id/convert-to-po",
        "GET /api/quotations/:id/pdf",
    ],
    validations=detail_validation(),
    acceptance=[
        "Trạng thái chuyển đúng workflow.",
        "PDF xuất đúng template brand.",
        "Audit log ghi đủ action + actor + thời gian.",
    ],
))


add(Screen(
    module="sales", module_label="Sales",
    slug="quotations-edit", title_vn="Chỉnh sửa báo giá", title_en="Edit Quotation",
    route="/sales/quotations/[id]/edit", kind="edit",
    purpose="Sửa báo giá khi còn ở trạng thái draft hoặc sent (nếu được phép).",
    actors=["Sales Staff", "Sales Manager"],
    roles=["sales", "sales-manager"],
    main_flow=[
        "Từ detail nhấn Sửa ⇒ load form với dữ liệu hiện tại.",
        "User thay đổi field/line items ⇒ nút Lưu bật.",
        "Lưu ⇒ `PATCH /api/quotations/:id` và quay về detail.",
    ],
    alt_flow=[
        "Báo giá đã `approved`/`converted` ⇒ redirect về detail, toast cảnh báo.",
        "Rời trang khi có thay đổi ⇒ dialog xác nhận.",
    ],
    layout=[
        "Reuse form component từ screen `Tạo báo giá`.",
        "Toolbar: Lưu, Huỷ.",
    ],
    fields=[
        Field("customerId", "uuid"), Field("validUntil", "date"),
        Field("items", "array"), Field("vatRate", "number"),
        Field("shippingTerms", "text"), Field("paymentTerms", "text"),
        Field("notes", "richtext"),
    ],
    actions=[
        "Lưu thay đổi.",
        "Huỷ bỏ.",
        "Thêm/xoá dòng items.",
    ],
    business_rules=[
        "Chỉ owner hoặc manager được sửa.",
        "Phiên bản được lưu vào history sau mỗi lần lưu.",
    ],
    apis=[
        "GET /api/quotations/:id",
        "PATCH /api/quotations/:id",
    ],
    validations=edit_validation(),
    acceptance=[
        "Không mất dữ liệu nếu điều hướng nhầm.",
        "Toast success khi lưu, fail khi validation.",
    ],
))


# -- Sales: Purchase Orders --------------------------------------------------

add(Screen(
    module="sales", module_label="Sales",
    slug="purchase-orders-list", title_vn="Danh sách đơn hàng mua", title_en="Purchase Orders List",
    route="/sales/purchase-orders", kind="list",
    purpose="Quản lý đơn hàng mua (PO) từ nhà cung cấp: tạo, duyệt, theo dõi giao hàng.",
    actors=["Purchasing Staff", "Purchasing Manager", "Admin"],
    roles=["purchasing", "purchasing-manager", "admin"],
    main_flow=[
        "Load `/sales/purchase-orders` ⇒ `GET /api/purchase-orders`.",
        "Filter theo status (draft, pending, approved, partial, received, cancelled).",
        "Click `+ Tạo PO` ⇒ form tạo.",
    ],
    alt_flow=[
        "Purchasing Staff chỉ thấy PO do mình tạo.",
    ],
    layout=[
        "Toolbar: search, filter status, supplier, ngày; + Tạo, Export.",
        "Body: DataGrid với cột mã, supplier, tổng tiền, status, expectedDeliveryDate.",
    ],
    fields=[
        Field("code", "string"),
        Field("supplierId", "uuid"),
        Field("totalAmount", "decimal"),
        Field("status", "enum(draft,pending,approved,partial,received,cancelled)"),
        Field("expectedDeliveryDate", "date"),
        Field("createdBy", "uuid"),
    ],
    actions=LIST_ACTIONS + ["Nút `Duyệt` bulk (chỉ manager)."],
    business_rules=[
        "Không cho xoá PO đã approved.",
        "PO từ quotation giữ reference quotationId.",
    ],
    apis=[
        "GET /api/purchase-orders?status=&supplierId=&q=",
        "POST /api/purchase-orders",
        "DELETE /api/purchase-orders/:id",
    ],
    validations=list_validation(),
    acceptance=[
        "Filter và search chính xác.",
        "Bulk approve chỉ hiện cho manager.",
    ],
))


add(Screen(
    module="sales", module_label="Sales",
    slug="purchase-orders-detail", title_vn="Chi tiết đơn hàng mua", title_en="Purchase Order Detail",
    route="/sales/purchase-orders/[id]", kind="detail",
    purpose="Xem chi tiết PO, tiến độ giao hàng, tạo phiếu nhập kho (Inbound) và gắn invoice/packing list.",
    actors=["Purchasing Staff", "Purchasing Manager", "Warehouse Staff"],
    roles=["purchasing", "warehouse", "admin"],
    main_flow=[
        "Click dòng PO ⇒ detail page.",
        "Hiển thị line items, tiến độ (received qty vs ordered qty).",
        "Nút `Tạo phiếu nhập kho` ⇒ tạo Inbound từ PO.",
        "Gắn invoice / packing list.",
    ],
    alt_flow=[
        "PO đã received full ⇒ disable Tạo phiếu nhập kho.",
    ],
    layout=[
        "Header: mã PO, supplier, status badge, nút hành động.",
        "Section line items + progress bar từng line.",
        "Section tài liệu: invoices, packing lists, inbound records.",
        "Section audit log.",
    ],
    fields=[
        Field("code", "string"), Field("supplierId", "uuid"),
        Field("items", "array"), Field("receivedItems", "array"),
        Field("status", "enum"), Field("expectedDeliveryDate", "date"),
        Field("invoices", "array<uuid>"), Field("packingLists", "array<uuid>"),
    ],
    actions=DETAIL_ACTIONS + [
        "Tạo phiếu nhập kho.",
        "Duyệt PO.",
        "Huỷ PO (chỉ manager).",
    ],
    business_rules=[
        "Nhập kho ≤ ordered qty của từng line.",
        "Cancel PO chỉ khi chưa có inbound.",
    ],
    apis=[
        "GET /api/purchase-orders/:id",
        "PATCH /api/purchase-orders/:id",
        "POST /api/purchase-orders/:id/approve",
        "POST /api/purchase-orders/:id/cancel",
        "POST /api/purchase-orders/:id/inbound",
    ],
    validations=detail_validation(),
    acceptance=[
        "Progress bar phản ánh đúng % received.",
        "Workflow status đúng thứ tự.",
        "Liên kết invoice/packing list chính xác.",
    ],
))


add(Screen(
    module="sales", module_label="Sales",
    slug="purchase-orders-edit", title_vn="Chỉnh sửa đơn hàng mua", title_en="Edit Purchase Order",
    route="/sales/purchase-orders/[id]/edit", kind="edit",
    purpose="Sửa PO khi còn ở trạng thái draft/pending.",
    actors=["Purchasing Staff", "Purchasing Manager"],
    roles=["purchasing", "purchasing-manager"],
    main_flow=[
        "Từ detail bấm Sửa ⇒ load form.",
        "User thay đổi line items / supplier / dates.",
        "Lưu ⇒ `PATCH /api/purchase-orders/:id` và về detail.",
    ],
    alt_flow=[
        "PO đã approved ⇒ redirect detail, toast lỗi.",
    ],
    layout=[
        "Reuse form từ `Tạo PO` (chưa có riêng nên xem là edit form).",
    ],
    fields=[
        Field("supplierId", "uuid"), Field("items", "array"),
        Field("expectedDeliveryDate", "date"), Field("notes", "richtext"),
        Field("shippingTerms", "text"), Field("paymentTerms", "text"),
    ],
    actions=["Lưu thay đổi", "Huỷ bỏ"],
    business_rules=[
        "Chỉ sửa khi status ∈ {draft, pending}.",
        "Mỗi lần sửa ghi audit.",
    ],
    apis=[
        "GET /api/purchase-orders/:id",
        "PATCH /api/purchase-orders/:id",
    ],
    validations=edit_validation(),
    acceptance=[
        "Không cho sửa khi status không hợp lệ.",
        "Audit log ghi lại thay đổi.",
    ],
))


# -- Warehouse ---------------------------------------------------------------

add(Screen(
    module="warehouse", module_label="Warehouse",
    slug="inbound-list", title_vn="Danh sách phiếu nhập kho", title_en="Inbound List",
    route="/warehouse/inbound", kind="list",
    purpose="Xem toàn bộ phiếu nhập kho: từ PO, từ trả hàng, hoặc điều chuyển.",
    actors=["Warehouse Staff", "Warehouse Manager"],
    roles=["warehouse", "warehouse-manager"],
    main_flow=[
        "Load `/warehouse/inbound` ⇒ `GET /api/inbounds`.",
        "Filter theo status (draft, posted), loại nguồn, warehouse, ngày.",
        "Tạo mới từ nút `+ Nhập kho`.",
    ],
    alt_flow=[
        "Warehouse Staff không được xoá phiếu posted.",
    ],
    layout=[
        "Toolbar: search, filter warehouse, source (PO/Return/Transfer), status, ngày; + Tạo.",
        "Body: DataGrid cột mã, source, warehouse, tổng số lượng, status, postedAt.",
    ],
    fields=[
        Field("code", "string"), Field("sourceType", "enum(po,return,transfer)"),
        Field("sourceId", "uuid"), Field("warehouseId", "uuid"),
        Field("totalQty", "number"), Field("status", "enum(draft,posted,cancelled)"),
        Field("postedAt", "datetime"),
    ],
    actions=LIST_ACTIONS,
    business_rules=[
        "Posted phiếu nhập kho cập nhật Inventory ngay.",
        "Không được sửa phiếu đã posted, chỉ tạo phiếu điều chỉnh.",
    ],
    apis=[
        "GET /api/inbounds?warehouseId=&sourceType=&status=",
        "POST /api/inbounds",
        "POST /api/inbounds/:id/post",
        "DELETE /api/inbounds/:id",
    ],
    validations=list_validation(),
    acceptance=[
        "Filter theo warehouse/source chính xác.",
        "Nút Post chỉ hiện khi draft.",
    ],
))


add(Screen(
    module="warehouse", module_label="Warehouse",
    slug="inbound-detail", title_vn="Chi tiết phiếu nhập kho", title_en="Inbound Detail",
    route="/warehouse/inbound/[id]", kind="detail",
    purpose="Xem chi tiết phiếu nhập kho và thực hiện post/cancel.",
    actors=["Warehouse Staff", "Warehouse Manager"],
    roles=["warehouse", "warehouse-manager"],
    main_flow=[
        "Click dòng inbound ⇒ detail.",
        "Hiển thị header (source, warehouse), lines (product, qty, lot, expiry).",
        "Nút Post (nếu draft), In phiếu.",
    ],
    alt_flow=[
        "Source là PO ⇒ hiển thị tham chiếu sang PO detail.",
        "Posted ⇒ chỉ xem.",
    ],
    layout=[
        "Header: mã, source link, warehouse, status.",
        "Section lines: product, qty, lot, expiry, note.",
        "Section audit.",
    ],
    fields=[
        Field("code", "string"), Field("sourceType", "enum"), Field("sourceId", "uuid"),
        Field("warehouseId", "uuid"),
        Field("lines", "array<{productId, qty, lot?, expiry?, note?}>"),
        Field("status", "enum"), Field("postedAt", "datetime"),
    ],
    actions=DETAIL_ACTIONS + [
        "Post (chuyển draft → posted, update inventory).",
        "Cancel (chỉ admin).",
        "In phiếu PDF.",
    ],
    business_rules=[
        "Post atomic; nếu tồn kho trùng lot đã có ⇒ cộng dồn.",
        "Cancel sau khi post ⇒ tạo phiếu điều chỉnh ngược, không xoá gốc.",
    ],
    apis=[
        "GET /api/inbounds/:id",
        "POST /api/inbounds/:id/post",
        "POST /api/inbounds/:id/cancel",
        "GET /api/inbounds/:id/pdf",
    ],
    validations=detail_validation(),
    acceptance=[
        "Post cập nhật đúng số lượng ở Inventory.",
        "Cancel tạo phiếu ngược.",
        "PDF chuẩn template.",
    ],
))


add(Screen(
    module="warehouse", module_label="Warehouse",
    slug="inventory-list", title_vn="Danh sách tồn kho", title_en="Inventory List",
    route="/warehouse/inventory", kind="list",
    purpose="Xem tồn kho realtime theo sản phẩm × kho × lot.",
    actors=["Warehouse Staff", "Warehouse Manager", "Sales Staff"],
    roles=["warehouse", "sales", "admin"],
    main_flow=[
        "Load `/warehouse/inventory` ⇒ `GET /api/inventory`.",
        "Filter theo warehouse, product, low-stock only.",
        "Hiển thị: product, warehouse, lot, availableQty, reservedQty, onHandQty.",
    ],
    alt_flow=[
        "Low-stock filter ⇒ chỉ hiện dòng có onHandQty <= reorderLevel.",
    ],
    layout=[
        "Toolbar: filter warehouse, product, toggle low-stock, Export.",
        "Body: DataGrid sortable.",
    ],
    fields=[
        Field("productId", "uuid"), Field("warehouseId", "uuid"),
        Field("lot", "string?"), Field("expiry", "date?"),
        Field("onHandQty", "number"), Field("reservedQty", "number"),
        Field("availableQty", "number", note="onHand - reserved"),
        Field("reorderLevel", "number"),
    ],
    actions=[
        "Export Excel theo filter hiện tại.",
        "Click dòng ⇒ chi tiết tồn kho.",
        "Không cho tạo mới (inventory derive từ inbound/outbound).",
    ],
    business_rules=[
        "availableQty = onHand - reserved.",
        "Reserved sinh từ quotation approved chưa ship.",
    ],
    apis=[
        "GET /api/inventory?warehouseId=&productId=&lowStock=",
        "GET /api/inventory/export.xlsx",
    ],
    validations=list_validation(),
    acceptance=[
        "Con số khớp với sum inbound - outbound.",
        "Low-stock filter hoạt động.",
    ],
))


add(Screen(
    module="warehouse", module_label="Warehouse",
    slug="inventory-detail", title_vn="Chi tiết tồn kho", title_en="Inventory Detail",
    route="/warehouse/inventory/[id]", kind="detail",
    purpose="Xem lịch sử giao dịch tồn kho của một product × warehouse × lot cụ thể.",
    actors=["Warehouse Staff", "Warehouse Manager"],
    roles=["warehouse", "warehouse-manager"],
    main_flow=[
        "Click dòng inventory ⇒ detail.",
        "Gọi `GET /api/inventory/:id` (+transactions).",
        "Hiển thị header (product, warehouse, lot) và table transactions (inbound/outbound/adjustment).",
    ],
    alt_flow=[
        "Tồn = 0 ⇒ hiển thị badge `Đã hết` và chỉ hiện lịch sử.",
    ],
    layout=[
        "Header: product, warehouse, lot, expiry, onHand, reserved, available.",
        "Section transactions: type, qty, balance, refDoc, at.",
    ],
    fields=[
        Field("productId", "uuid"), Field("warehouseId", "uuid"),
        Field("lot", "string"), Field("expiry", "date?"),
        Field("onHandQty", "number"), Field("reservedQty", "number"),
        Field("transactions", "array<{type,qty,balance,refDocId,at}>"),
    ],
    actions=[
        "Export lịch sử.",
        "Điều chỉnh thủ công (chỉ manager) ⇒ drawer adjustment.",
        "Link sang phiếu refDoc (inbound/outbound).",
    ],
    business_rules=[
        "Điều chỉnh tạo record transaction type=adjustment.",
        "Không cho edit transaction lịch sử.",
    ],
    apis=[
        "GET /api/inventory/:id",
        "GET /api/inventory/:id/transactions?page=",
        "POST /api/inventory/:id/adjust",
    ],
    validations=detail_validation(),
    acceptance=[
        "Balance cuối cùng khớp onHandQty.",
        "Điều chỉnh có lý do bắt buộc; ghi audit.",
    ],
))


add(Screen(
    module="warehouse", module_label="Warehouse",
    slug="outbound-list", title_vn="Danh sách phiếu xuất kho", title_en="Outbound List",
    route="/warehouse/outbound", kind="list",
    purpose="Quản lý phiếu xuất kho: giao hàng khách, chuyển kho, trả nhà cung cấp.",
    actors=["Warehouse Staff", "Warehouse Manager", "Sales Staff"],
    roles=["warehouse", "sales", "admin"],
    main_flow=[
        "Load `/warehouse/outbound` ⇒ `GET /api/outbounds`.",
        "Filter theo status (draft, allocated, picked, shipped), target (customer, warehouse, supplier), ngày.",
        "Tạo mới từ nút `+ Xuất kho` (từ PO bán / quotation / thủ công).",
    ],
    alt_flow=[
        "Không đủ tồn kho allocate ⇒ hiển thị cảnh báo, giữ ở draft.",
    ],
    layout=[
        "Toolbar: filter warehouse, target type, status; + Tạo.",
        "Body: DataGrid cột mã, target, warehouse, tổng qty, status, scheduledAt.",
    ],
    fields=[
        Field("code", "string"),
        Field("targetType", "enum(customer, warehouse, supplier)"),
        Field("targetId", "uuid"),
        Field("warehouseId", "uuid"),
        Field("totalQty", "number"),
        Field("status", "enum(draft,allocated,picked,shipped,cancelled)"),
        Field("scheduledAt", "datetime"),
    ],
    actions=LIST_ACTIONS + [
        "Chuyển trạng thái theo workflow (allocate, pick, ship).",
        "Huỷ phiếu nếu draft.",
    ],
    business_rules=[
        "Allocate giữ reservedQty; Ship trừ onHandQty.",
        "Không cho xoá phiếu đã ship.",
    ],
    apis=[
        "GET /api/outbounds?status=&warehouseId=&targetType=",
        "POST /api/outbounds",
        "POST /api/outbounds/:id/allocate",
        "POST /api/outbounds/:id/ship",
        "DELETE /api/outbounds/:id",
    ],
    validations=list_validation(),
    acceptance=[
        "Các trạng thái chuyển đúng workflow.",
        "Reserved & onHand cập nhật đúng sau mỗi action.",
    ],
))


add(Screen(
    module="warehouse", module_label="Warehouse",
    slug="outbound-detail", title_vn="Chi tiết phiếu xuất kho", title_en="Outbound Detail",
    route="/warehouse/outbound/[id]", kind="detail",
    purpose="Xem chi tiết phiếu xuất kho, allocate tồn kho, tạo packing list, in phiếu giao hàng.",
    actors=["Warehouse Staff", "Warehouse Manager"],
    roles=["warehouse", "warehouse-manager"],
    main_flow=[
        "Click dòng outbound ⇒ detail.",
        "Hiển thị header (target, warehouse, status), lines (product, qty, allocated lots).",
        "Nút Allocate/Pick/Ship/Cancel.",
        "Liên kết sang packing list.",
    ],
    alt_flow=[
        "Shipped ⇒ disable mọi action trừ in phiếu.",
    ],
    layout=[
        "Header: mã, target, warehouse, status badge.",
        "Section lines: product, requested qty, allocated qty, lot/expiry.",
        "Section tài liệu: packing list liên kết, invoice (nếu có).",
        "Audit log.",
    ],
    fields=[
        Field("code", "string"), Field("targetType", "enum"), Field("targetId", "uuid"),
        Field("warehouseId", "uuid"),
        Field("lines", "array<{productId, requestedQty, allocatedQty, lots:[{lot,qty}]}>"),
        Field("packingListId", "uuid?"),
        Field("status", "enum"),
    ],
    actions=DETAIL_ACTIONS + [
        "Allocate / Pick / Ship / Cancel.",
        "Tạo packing list.",
        "In phiếu giao hàng PDF.",
    ],
    business_rules=[
        "Ship atomic: cập nhật inventory + tạo transaction.",
        "Cancel sau ship ⇒ tạo phiếu nhập ngược (return).",
    ],
    apis=[
        "GET /api/outbounds/:id",
        "POST /api/outbounds/:id/allocate",
        "POST /api/outbounds/:id/pick",
        "POST /api/outbounds/:id/ship",
        "POST /api/outbounds/:id/cancel",
        "GET /api/outbounds/:id/pdf",
    ],
    validations=detail_validation(),
    acceptance=[
        "allocatedQty <= requestedQty <= onHandQty.",
        "Workflow không skip bước.",
        "PDF phiếu giao hàng đúng chuẩn.",
    ],
))


add(Screen(
    module="warehouse", module_label="Warehouse",
    slug="outbound-inventory-list", title_vn="Danh sách tồn kho xuất", title_en="Outbound Inventory List",
    route="/warehouse/outbound/inventory", kind="list",
    purpose="Xem tồn kho đã được allocate/reserved cho các phiếu xuất nhưng chưa shipped.",
    actors=["Warehouse Staff", "Warehouse Manager"],
    roles=["warehouse", "warehouse-manager"],
    main_flow=[
        "Load `/warehouse/outbound/inventory` ⇒ `GET /api/outbound-inventory`.",
        "Hiển thị: product, warehouse, lot, reservedQty, outboundCode liên kết.",
        "Filter theo warehouse, product, outboundId.",
    ],
    alt_flow=[
        "Outbound bị cancel ⇒ dòng không hiển thị.",
    ],
    layout=[
        "Toolbar: filter warehouse, product, outbound code; Export.",
        "Body: DataGrid.",
    ],
    fields=[
        Field("productId", "uuid"), Field("warehouseId", "uuid"),
        Field("lot", "string?"), Field("reservedQty", "number"),
        Field("outboundId", "uuid"), Field("outboundCode", "string"),
        Field("status", "enum(allocated,picked)"),
    ],
    actions=[
        "Click link outboundCode ⇒ `/warehouse/outbound/[id]`.",
        "Export Excel theo filter.",
    ],
    business_rules=[
        "Dữ liệu derive từ outbound × lines với status allocated/picked.",
    ],
    apis=[
        "GET /api/outbound-inventory?warehouseId=&productId=",
    ],
    validations=list_validation(),
    acceptance=[
        "Con số reserved khớp `/warehouse/inventory`.",
        "Export hoạt động.",
    ],
))


add(Screen(
    module="warehouse", module_label="Warehouse",
    slug="outbound-inventory-detail", title_vn="Chi tiết tồn kho xuất", title_en="Outbound Inventory Detail",
    route="/warehouse/outbound/inventory/[id]", kind="detail",
    purpose="Xem chi tiết một bản ghi tồn kho xuất: liên kết outbound, lịch sử allocate.",
    actors=["Warehouse Staff", "Warehouse Manager"],
    roles=["warehouse", "warehouse-manager"],
    main_flow=[
        "Click dòng outbound-inventory ⇒ detail.",
        "Hiển thị product, lot, reservedQty, lịch sử allocate/pick.",
    ],
    alt_flow=[
        "Nếu outbound đã ship ⇒ chuyển sang trang transaction history.",
    ],
    layout=[
        "Header: product, warehouse, lot, outboundCode.",
        "Section timeline: allocate → pick → ship.",
    ],
    fields=[
        Field("productId", "uuid"), Field("warehouseId", "uuid"),
        Field("lot", "string"), Field("reservedQty", "number"),
        Field("outboundId", "uuid"),
        Field("events", "array<{type, qty, at, by}>"),
    ],
    actions=["Link sang outbound detail", "Export lịch sử"],
    business_rules=[
        "Không cho chỉnh sửa trực tiếp, mọi thay đổi phải qua outbound.",
    ],
    apis=[
        "GET /api/outbound-inventory/:id",
        "GET /api/outbound-inventory/:id/events",
    ],
    validations=detail_validation(),
    acceptance=[
        "Timeline đủ events.",
        "Link về outbound chính xác.",
    ],
))


# -- Logistics ---------------------------------------------------------------

add(Screen(
    module="logistics", module_label="Logistics",
    slug="container-planning-list", title_vn="Danh sách kế hoạch container", title_en="Container Planning List",
    route="/logistics/container-planning", kind="list",
    purpose="Lập và theo dõi kế hoạch đóng hàng vào container cho các outbound.",
    actors=["Logistics Staff", "Logistics Manager"],
    roles=["logistics", "logistics-manager"],
    main_flow=[
        "Load `/logistics/container-planning` ⇒ `GET /api/container-plans`.",
        "Filter theo status (draft, confirmed, cancelled), ngày dự kiến.",
        "Click `+ Tạo kế hoạch` để tạo mới.",
    ],
    alt_flow=[
        "Logistics Staff chỉ thấy plan của mình.",
    ],
    layout=[
        "Toolbar: filter status, ngày, search; + Tạo.",
        "Body: DataGrid cột mã plan, container type, status, ngày dự kiến, utilization%.",
    ],
    fields=[
        Field("code", "string"), Field("containerType", "enum(20ft,40ft,40hc)"),
        Field("plannedDate", "date"),
        Field("status", "enum(draft,confirmed,cancelled)"),
        Field("utilization", "number", note="0-100%"),
    ],
    actions=LIST_ACTIONS,
    business_rules=[
        "Utilization tính dựa trên thể tích lines / dung tích container.",
        "Không cho xoá plan đã confirmed (phải cancel).",
    ],
    apis=[
        "GET /api/container-plans?status=&date=",
        "POST /api/container-plans",
        "DELETE /api/container-plans/:id",
    ],
    validations=list_validation(),
    acceptance=[
        "Utilization chính xác theo ccfg container.",
        "Filter status hoạt động.",
    ],
))


add(Screen(
    module="logistics", module_label="Logistics",
    slug="container-planning-detail", title_vn="Chi tiết kế hoạch container", title_en="Container Plan Detail",
    route="/logistics/container-planning/[id]", kind="detail",
    purpose="Xem và chỉnh sửa kế hoạch đóng container: thêm outbound vào plan, sắp xếp thứ tự.",
    actors=["Logistics Staff", "Logistics Manager"],
    roles=["logistics", "logistics-manager"],
    main_flow=[
        "Click dòng plan ⇒ detail.",
        "Hiển thị header + danh sách outbound trong plan + visual 3D/2D đơn giản về utilization.",
        "Drag-drop thứ tự; thêm outbound còn trống.",
        "Xác nhận plan ⇒ chuyển status=confirmed.",
    ],
    alt_flow=[
        "Plan confirmed ⇒ disable chỉnh sửa, chỉ xem.",
    ],
    layout=[
        "Header: mã plan, container type, ngày, utilization.",
        "Section lines: list outbound trong plan.",
        "Section visualization: thanh progress + layout sơ đồ.",
    ],
    fields=[
        Field("code", "string"), Field("containerType", "enum"),
        Field("plannedDate", "date"),
        Field("lines", "array<{outboundId, order, volume, weight}>"),
        Field("utilization", "number"), Field("status", "enum"),
    ],
    actions=DETAIL_ACTIONS + [
        "Thêm outbound (chọn từ danh sách outbound shipped-ready).",
        "Xoá outbound khỏi plan.",
        "Confirm plan.",
        "Export PDF kế hoạch.",
    ],
    business_rules=[
        "Không cho thêm outbound đã thuộc plan khác.",
        "Utilization >100% ⇒ chặn confirm.",
    ],
    apis=[
        "GET /api/container-plans/:id",
        "PATCH /api/container-plans/:id",
        "POST /api/container-plans/:id/confirm",
        "POST /api/container-plans/:id/cancel",
        "POST /api/container-plans/:id/lines",
        "DELETE /api/container-plans/:id/lines/:lineId",
    ],
    validations=detail_validation(),
    acceptance=[
        "Drag-drop cập nhật order.",
        "Utilization bar đúng thực tế.",
        "Confirm hoạt động và khoá chỉnh sửa.",
    ],
))


add(Screen(
    module="logistics", module_label="Logistics",
    slug="container-loading-list", title_vn="Danh sách xếp container", title_en="Container Loading List",
    route="/logistics/container-loading", kind="list",
    purpose="Ghi nhận quá trình xếp container thực tế từ plan.",
    actors=["Logistics Staff", "Warehouse Staff"],
    roles=["logistics", "warehouse"],
    main_flow=[
        "Load `/logistics/container-loading` ⇒ `GET /api/container-loadings`.",
        "Filter theo status (open, in_progress, completed), ngày.",
        "Tạo mới từ plan confirmed.",
    ],
    alt_flow=[
        "Plan chưa confirm ⇒ không cho tạo loading.",
    ],
    layout=[
        "Toolbar: filter status, ngày; + Tạo từ plan.",
        "Body: DataGrid cột mã, plan liên kết, container number, status, startedAt, completedAt.",
    ],
    fields=[
        Field("code", "string"), Field("planId", "uuid", required=True),
        Field("containerNumber", "string"),
        Field("status", "enum(open,in_progress,completed,cancelled)"),
        Field("startedAt", "datetime"), Field("completedAt", "datetime"),
    ],
    actions=LIST_ACTIONS + [
        "Start / Complete loading.",
    ],
    business_rules=[
        "Completed ⇒ chuyển trạng thái outbound tương ứng sang shipped.",
    ],
    apis=[
        "GET /api/container-loadings?status=&date=",
        "POST /api/container-loadings",
        "POST /api/container-loadings/:id/start",
        "POST /api/container-loadings/:id/complete",
    ],
    validations=list_validation(),
    acceptance=[
        "Workflow status chuyển đúng.",
        "Completed cập nhật outbound shipped.",
    ],
))


add(Screen(
    module="logistics", module_label="Logistics",
    slug="container-loading-detail", title_vn="Chi tiết xếp container", title_en="Container Loading Detail",
    route="/logistics/container-loading/[id]", kind="detail",
    purpose="Ghi nhận chi tiết xếp hàng: lines đã load, hình ảnh, seal, nhiệt độ.",
    actors=["Logistics Staff", "Warehouse Staff"],
    roles=["logistics", "warehouse"],
    main_flow=[
        "Click dòng loading ⇒ detail.",
        "Tick từng line đã load, nhập số seal, upload hình.",
        "Complete khi đủ lines.",
    ],
    alt_flow=[
        "Lines thiếu so với plan ⇒ cảnh báo nhưng cho phép Complete có ghi chú.",
    ],
    layout=[
        "Header: mã loading, plan, container number, status.",
        "Section lines: checkbox loaded, ghi chú.",
        "Section photos & seals.",
    ],
    fields=[
        Field("code", "string"), Field("planId", "uuid"),
        Field("containerNumber", "string"), Field("sealNumbers", "array<string>"),
        Field("temperature", "number?"),
        Field("lines", "array<{planLineId, loaded, notes}>"),
        Field("photos", "array<url>"), Field("status", "enum"),
    ],
    actions=DETAIL_ACTIONS + [
        "Upload hình ảnh (drag-drop).",
        "Complete loading.",
        "In phiếu xếp container.",
    ],
    business_rules=[
        "Hình ảnh tối đa 20, mỗi file <= 5MB.",
        "Seal number bắt buộc khi Complete.",
    ],
    apis=[
        "GET /api/container-loadings/:id",
        "PATCH /api/container-loadings/:id",
        "POST /api/container-loadings/:id/complete",
        "POST /api/container-loadings/:id/photos",
    ],
    validations=detail_validation(),
    acceptance=[
        "Không cho Complete khi thiếu seal.",
        "Upload ảnh hoạt động với progress.",
    ],
))


add(Screen(
    module="logistics", module_label="Logistics",
    slug="logistics-overview", title_vn="Tổng quan logistics", title_en="Logistics Overview",
    route="/logistics", kind="other",
    purpose="Dashboard tổng quan logistics: số plan đang mở, loading đang chạy, container hôm nay.",
    actors=["Logistics Manager", "Admin"],
    roles=["logistics-manager", "admin"],
    main_flow=[
        "Load `/logistics` ⇒ gọi API KPI logistics.",
        "Hiển thị card: plan mở, loading đang chạy, container hoàn thành hôm nay.",
        "Bảng alerts: container utilization <60%, delay plan.",
    ],
    alt_flow=[
        "Không có dữ liệu ⇒ hiển thị empty state.",
    ],
    layout=[
        "Hàng 1: 3 KPI card.",
        "Hàng 2: chart daily loading.",
        "Hàng 3: table alerts.",
    ],
    fields=[
        Field("openPlans", "number"), Field("inProgressLoadings", "number"),
        Field("completedToday", "number"),
        Field("alerts", "array<{type, message, refId}>"),
    ],
    actions=[
        "Click KPI ⇒ filter sang list tương ứng.",
        "Click alert ⇒ sang detail liên quan.",
    ],
    business_rules=[
        "Alert utilization dưới ngưỡng config; delay nếu plannedDate < today và chưa complete.",
    ],
    apis=[
        "GET /api/logistics/overview",
        "GET /api/logistics/alerts",
    ],
    validations=list_validation(),
    acceptance=[
        "KPI khớp data list thực tế.",
        "Alerts có link điều hướng đúng.",
    ],
))


# -- Documents ---------------------------------------------------------------

add(Screen(
    module="documents", module_label="Documents",
    slug="invoices-list", title_vn="Danh sách hoá đơn", title_en="Invoices List",
    route="/documents/invoices", kind="list",
    purpose="Quản lý hoá đơn xuất cho khách hàng và hoá đơn nhận từ nhà cung cấp.",
    actors=["Accountant", "Admin"],
    roles=["accountant", "admin"],
    main_flow=[
        "Load `/documents/invoices` ⇒ `GET /api/invoices`.",
        "Filter theo type (sales/purchase), status, ngày.",
        "Click `+ Tạo hoá đơn` ⇒ drawer tạo.",
    ],
    alt_flow=[
        "Kế toán không thấy hoá đơn thuộc công ty con khác nếu đa-tenant.",
    ],
    layout=[
        "Toolbar: filter type, status, ngày, search; + Tạo, Export.",
        "Body: DataGrid cột mã, đối tác, totalAmount, status, issueDate.",
    ],
    fields=[
        Field("code", "string"), Field("type", "enum(sales,purchase)"),
        Field("partnerId", "uuid"), Field("totalAmount", "decimal"),
        Field("currency", "enum"), Field("issueDate", "date"),
        Field("status", "enum(draft,issued,paid,cancelled)"),
    ],
    actions=LIST_ACTIONS + [
        "Export Excel theo filter.",
        "In PDF hoá đơn.",
    ],
    business_rules=[
        "Invoice sales gắn với outbound/quotation.",
        "Invoice purchase gắn với PO/inbound.",
        "Không cho xoá invoice issued — chỉ cancel.",
    ],
    apis=[
        "GET /api/invoices?type=&status=&q=",
        "POST /api/invoices",
        "POST /api/invoices/:id/issue",
        "POST /api/invoices/:id/cancel",
    ],
    validations=list_validation(),
    acceptance=[
        "Filter đúng type.",
        "Export Excel đủ cột quy định kế toán.",
    ],
))


add(Screen(
    module="documents", module_label="Documents",
    slug="invoices-detail", title_vn="Chi tiết hoá đơn", title_en="Invoice Detail",
    route="/documents/invoices/[id]", kind="detail",
    purpose="Xem chi tiết hoá đơn, in PDF, đánh dấu đã thanh toán.",
    actors=["Accountant", "Admin"],
    roles=["accountant", "admin"],
    main_flow=[
        "Click dòng ⇒ detail invoice.",
        "Hiển thị header + items + thuế + tổng + thông tin ngân hàng.",
        "Nút Issue, Mark Paid, Cancel, In PDF.",
    ],
    alt_flow=[
        "Đã paid ⇒ disable Mark Paid.",
    ],
    layout=[
        "Header: mã, partner, type, status.",
        "Section items: product/description, qty, price, vat.",
        "Section totals.",
        "Section ghi chú & thanh toán.",
    ],
    fields=[
        Field("code", "string"), Field("type", "enum"),
        Field("partnerId", "uuid"), Field("items", "array"),
        Field("vatAmount", "decimal"), Field("totalAmount", "decimal"),
        Field("paymentMethod", "string"), Field("paidAt", "datetime"),
        Field("status", "enum"),
    ],
    actions=DETAIL_ACTIONS + ["Mark Paid", "Cancel invoice", "In PDF"],
    business_rules=[
        "Mark Paid cần ngày thanh toán + phương thức.",
        "Cancel invoice tạo debit note tham chiếu.",
    ],
    apis=[
        "GET /api/invoices/:id",
        "POST /api/invoices/:id/mark-paid",
        "POST /api/invoices/:id/cancel",
        "GET /api/invoices/:id/pdf",
    ],
    validations=detail_validation(),
    acceptance=[
        "PDF đúng mẫu thuế.",
        "Audit log đầy đủ.",
    ],
))


add(Screen(
    module="documents", module_label="Documents",
    slug="packing-lists-list", title_vn="Danh sách phiếu đóng gói", title_en="Packing Lists List",
    route="/documents/packing-lists", kind="list",
    purpose="Quản lý phiếu đóng gói gắn với outbound/shipment.",
    actors=["Warehouse Staff", "Logistics Staff"],
    roles=["warehouse", "logistics"],
    main_flow=[
        "Load `/documents/packing-lists` ⇒ `GET /api/packing-lists`.",
        "Filter theo outbound, status, ngày.",
        "Click `+ Tạo từ outbound` ⇒ chọn outbound ready.",
    ],
    alt_flow=[
        "Outbound chưa allocate ⇒ không thấy trong dropdown.",
    ],
    layout=[
        "Toolbar: filter status, ngày, search; + Tạo.",
        "Body: DataGrid cột mã, outbound, khách hàng, totalPackages, status.",
    ],
    fields=[
        Field("code", "string"), Field("outboundId", "uuid"),
        Field("customerId", "uuid"), Field("totalPackages", "number"),
        Field("totalWeight", "number"),
        Field("status", "enum(draft,issued,shipped,cancelled)"),
    ],
    actions=LIST_ACTIONS + ["In PDF packing list"],
    business_rules=[
        "Không cho xoá packing list đã shipped.",
    ],
    apis=[
        "GET /api/packing-lists",
        "POST /api/packing-lists",
        "DELETE /api/packing-lists/:id",
    ],
    validations=list_validation(),
    acceptance=[
        "Liên kết outbound chính xác.",
        "PDF chuẩn template.",
    ],
))


add(Screen(
    module="documents", module_label="Documents",
    slug="packing-lists-detail", title_vn="Chi tiết phiếu đóng gói", title_en="Packing List Detail",
    route="/documents/packing-lists/[id]", kind="detail",
    purpose="Xem chi tiết, chỉnh sửa các gói (package) và in PDF.",
    actors=["Warehouse Staff", "Logistics Staff"],
    roles=["warehouse", "logistics"],
    main_flow=[
        "Click dòng packing list ⇒ detail.",
        "Hiển thị header + list packages (carton): trọng lượng, kích thước, items bên trong.",
        "Nút Issue, In PDF, Ship.",
    ],
    alt_flow=[
        "Shipped ⇒ khoá chỉnh sửa.",
    ],
    layout=[
        "Header: mã, outbound, customer, status.",
        "Section packages: list carton + items.",
        "Section totals.",
    ],
    fields=[
        Field("code", "string"), Field("outboundId", "uuid"),
        Field("customerId", "uuid"),
        Field("packages", "array<{code, weight, dim, items:[{productId, qty}]}>"),
        Field("totalPackages", "number"), Field("totalWeight", "number"),
        Field("status", "enum"),
    ],
    actions=DETAIL_ACTIONS + [
        "Thêm package mới.",
        "Issue (xuất phiếu chính thức).",
        "In PDF.",
    ],
    business_rules=[
        "Tổng qty của packages phải khớp outbound lines (allow tolerance 0).",
    ],
    apis=[
        "GET /api/packing-lists/:id",
        "PATCH /api/packing-lists/:id",
        "POST /api/packing-lists/:id/issue",
        "GET /api/packing-lists/:id/pdf",
    ],
    validations=detail_validation(),
    acceptance=[
        "Tổng khớp outbound.",
        "PDF đủ cột HS code (nếu export).",
    ],
))


# -- System ------------------------------------------------------------------

add(Screen(
    module="system", module_label="System",
    slug="users-list", title_vn="Danh sách người dùng", title_en="Users List",
    route="/system/users", kind="list",
    purpose="Quản lý người dùng: tạo, gán role, reset mật khẩu, disable.",
    actors=["Admin", "IT"],
    roles=["admin"],
    main_flow=[
        "Load `/system/users` ⇒ `GET /api/users`.",
        "Filter theo role, status, department; search theo email/tên.",
        "Click `+ Tạo user` ⇒ drawer.",
    ],
    alt_flow=[
        "Không phải admin ⇒ trang 403.",
    ],
    layout=[
        "Toolbar: filter role, status, department; + Tạo, Import.",
        "Body: DataGrid cột email, name, role, department, status, lastLogin.",
    ],
    fields=[
        Field("email", "email", required=True),
        Field("name", "string", required=True),
        Field("role", "enum"), Field("department", "string"),
        Field("status", "enum(active,inactive,locked)"),
        Field("lastLoginAt", "datetime"),
    ],
    actions=LIST_ACTIONS + [
        "Reset password (gửi mail token).",
        "Disable / enable user.",
    ],
    business_rules=[
        "Chỉ admin được thao tác.",
        "Không được disable chính mình.",
    ],
    apis=[
        "GET /api/users?role=&status=",
        "POST /api/users",
        "PATCH /api/users/:id",
        "POST /api/users/:id/reset-password",
    ],
    validations=list_validation() + [
        "Email duy nhất, đúng định dạng.",
    ],
    acceptance=[
        "Tạo user gửi mail invite.",
        "Role viewer không thấy trang.",
    ],
))


add(Screen(
    module="system", module_label="System",
    slug="users-import", title_vn="Import người dùng", title_en="Users Import",
    route="/system/users/import", kind="other",
    purpose="Import hàng loạt users từ file CSV/Excel với preview và xử lý lỗi.",
    actors=["Admin"],
    roles=["admin"],
    main_flow=[
        "User bấm Import ⇒ vào `/system/users/import`.",
        "Upload file Excel/CSV (max 5MB).",
        "Server validate, trả preview bảng có trạng thái từng dòng.",
        "User confirm ⇒ import thật.",
    ],
    alt_flow=[
        "File sai định dạng ⇒ hiển thị lỗi chi tiết, cho tải template.",
        "Dòng lỗi validate ⇒ cho phép fix nhanh inline hoặc skip.",
    ],
    layout=[
        "Bước 1: upload file + nút tải template.",
        "Bước 2: preview bảng + bộ lọc lỗi/ok.",
        "Bước 3: confirm + progress bar + summary.",
    ],
    fields=[
        Field("file", "file", required=True),
        Field("skipInvalid", "boolean", note="checkbox trước khi confirm"),
    ],
    actions=[
        "Tải file template.",
        "Upload file.",
        "Confirm import.",
        "Huỷ.",
    ],
    business_rules=[
        "Email trùng ⇒ skip hoặc overwrite tuỳ config.",
        "Tối đa 10,000 dòng mỗi lần import.",
    ],
    apis=[
        "GET /api/users/import-template.xlsx",
        "POST /api/users/import/preview (multipart)",
        "POST /api/users/import/confirm",
    ],
    validations=[
        "File size <= 5MB; ext ∈ {xlsx, csv}.",
        "Các trường bắt buộc đủ; email đúng định dạng.",
        "Progress bar update theo batch 200 dòng.",
    ],
    acceptance=[
        "Import 5000 rows <30s, có progress.",
        "Báo cáo summary: ok / skip / fail.",
        "Có log chi tiết tải xuống.",
    ],
))


add(Screen(
    module="system", module_label="System",
    slug="settings", title_vn="Cài đặt hệ thống", title_en="System Settings",
    route="/system/settings", kind="other",
    purpose="Cấu hình chung: thông tin công ty, múi giờ, tiền tệ, email SMTP, cấu hình kho mặc định, tham số business.",
    actors=["Admin"],
    roles=["admin"],
    main_flow=[
        "Load `/system/settings` ⇒ `GET /api/settings`.",
        "Render form dạng tab: Company, Localization, Email, Warehouse, Business Rules.",
        "Lưu ⇒ `PATCH /api/settings` per tab.",
    ],
    alt_flow=[
        "Kết nối SMTP fail ⇒ hiển thị lỗi và giữ config chưa save.",
    ],
    layout=[
        "Sidebar tab trái, form bên phải.",
        "Footer: nút Lưu / Huỷ / Test.",
    ],
    fields=[
        Field("companyName", "string"), Field("taxCode", "string"),
        Field("timezone", "enum", note="vd Asia/Ho_Chi_Minh"),
        Field("defaultCurrency", "enum"), Field("dateFormat", "string"),
        Field("smtpHost", "string"), Field("smtpPort", "number"),
        Field("smtpUser", "string"), Field("smtpPassword", "password"),
        Field("defaultWarehouseId", "uuid"), Field("vatDefault", "number"),
        Field("reorderLevelDefault", "number"),
    ],
    actions=[
        "Lưu từng tab.",
        "Test SMTP (gửi mail test).",
        "Revert thay đổi.",
    ],
    business_rules=[
        "Chỉ admin được truy cập.",
        "Đổi timezone cập nhật mọi hiển thị datetime.",
    ],
    apis=[
        "GET /api/settings",
        "PATCH /api/settings",
        "POST /api/settings/test-smtp",
    ],
    validations=create_validation() + [
        "Port ∈ [1, 65535].",
        "Email test phải đúng format.",
    ],
    acceptance=[
        "Test SMTP trả kết quả rõ ràng.",
        "Lưu thành công ⇒ toast và cập nhật state.",
    ],
))


# -- Notifications -----------------------------------------------------------

add(Screen(
    module="notifications", module_label="Notifications",
    slug="notifications-list", title_vn="Danh sách thông báo", title_en="Notifications List",
    route="/notifications", kind="list",
    purpose="Hiển thị toàn bộ thông báo hệ thống gửi tới user: low-stock, duyệt PO, shipment delay...",
    actors=["All authenticated users"],
    roles=["*"],
    main_flow=[
        "Load `/notifications` ⇒ `GET /api/notifications?userId=me`.",
        "Filter theo loại (alert, info, task), trạng thái (unread/read).",
        "Click dòng ⇒ mark read + điều hướng theo `link`.",
    ],
    alt_flow=[
        "Realtime: socket.io push notification mới ⇒ badge header cập nhật.",
    ],
    layout=[
        "Toolbar: filter type, status, ngày; nút Mark all read.",
        "Body: DataGrid cột icon, title, body, createdAt, actions.",
    ],
    fields=[
        Field("title", "string"), Field("body", "text"),
        Field("type", "enum(alert,info,task)"),
        Field("link", "string", note="route client điều hướng"),
        Field("readAt", "datetime?"),
        Field("createdAt", "datetime"),
    ],
    actions=[
        "Mark all read.",
        "Mark individual read/unread.",
        "Click điều hướng theo link.",
    ],
    business_rules=[
        "Mỗi notification thuộc user cụ thể.",
        "Giữ 90 ngày rồi purge (background job).",
    ],
    apis=[
        "GET /api/notifications?userId=me&type=&status=",
        "PATCH /api/notifications/:id/read",
        "POST /api/notifications/mark-all-read",
    ],
    validations=list_validation(),
    acceptance=[
        "Realtime cập nhật không cần reload.",
        "Filter unread chính xác.",
    ],
))


# ---------------------------------------------------------------------------
# Rendering
# ---------------------------------------------------------------------------


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_kv(doc: Document, pairs: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for k, v in pairs:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v


def add_fields_table(doc: Document, fields: list[Field]) -> None:
    if not fields:
        return
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Type"
    hdr[2].text = "Required"
    hdr[3].text = "Ghi chú"
    for f in fields:
        cells = table.add_row().cells
        cells[0].text = f.name
        cells[1].text = f.type
        cells[2].text = "Yes" if f.required else "No"
        cells[3].text = f.note


def render_screen(screen: Screen) -> Path:
    out_dir = OUTPUT_ROOT / screen.module
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"{screen.slug}.docx"

    doc = Document()

    # Title
    title = doc.add_heading(f"Basic Design — {screen.title_vn}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph()
    run = subtitle.add_run(f"({screen.title_en})")
    run.italic = True

    add_heading(doc, "1. Thông tin chung", level=1)
    add_kv(doc, [
        ("Module", screen.module_label),
        ("Route", screen.route),
        ("Kind", screen.kind),
        ("Screen ID", f"{screen.module}/{screen.slug}"),
        ("Project", "Inventory Management System Development"),
        ("Ticket", "DXSA-10"),
    ])

    add_heading(doc, "2. Mục đích & phạm vi", level=1)
    doc.add_paragraph(screen.purpose)

    add_heading(doc, "3. Actors & quyền truy cập (RBAC)", level=1)
    doc.add_paragraph("Actors:")
    add_bullets(doc, screen.actors)
    doc.add_paragraph("Role được phép:")
    add_bullets(doc, screen.roles if screen.roles else ["(theo mặc định hệ thống)"])

    add_heading(doc, "4. Luồng nghiệp vụ", level=1)
    doc.add_paragraph("Luồng chính:")
    add_numbered(doc, screen.main_flow)
    if screen.alt_flow:
        doc.add_paragraph("Luồng phụ / ngoại lệ:")
        add_bullets(doc, screen.alt_flow)

    add_heading(doc, "5. Layout & wireframe (mô tả)", level=1)
    add_bullets(doc, screen.layout)

    add_heading(doc, "6. Thành phần dữ liệu (fields)", level=1)
    add_fields_table(doc, screen.fields)

    add_heading(doc, "7. Hành động (actions)", level=1)
    add_bullets(doc, screen.actions)

    add_heading(doc, "8. Business rules", level=1)
    add_bullets(doc, screen.business_rules)

    add_heading(doc, "9. API endpoints", level=1)
    add_bullets(doc, screen.apis)

    add_heading(doc, "10. Validation & Error handling", level=1)
    add_bullets(doc, screen.validations)
    doc.add_paragraph("Chuẩn error handling chung:")
    add_bullets(doc, COMMON_ERROR_RULES)

    add_heading(doc, "11. Acceptance criteria", level=1)
    add_bullets(doc, screen.acceptance)

    add_heading(doc, "12. Open questions", level=1)
    if screen.open_questions:
        add_bullets(doc, screen.open_questions)
    else:
        doc.add_paragraph("(chưa có)")

    add_heading(doc, "Footer", level=1)
    doc.add_paragraph(
        "Tài liệu được tạo tự động bởi Product Manager agent theo ticket DXSA-10. "
        "Cập nhật qua script `scripts/generate_screen_design_docs.py`."
    )

    # Compact spacing
    for p in doc.paragraphs:
        if p.style.name == "Normal":
            p.paragraph_format.space_after = Pt(4)

    doc.save(path)
    return path


def main() -> None:
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    generated: list[Path] = []
    for screen in SCREENS:
        path = render_screen(screen)
        generated.append(path)
    print(f"Generated {len(generated)} docx files under {OUTPUT_ROOT}")
    for p in generated:
        print(f"  - {p.relative_to(REPO_ROOT)}")


if __name__ == "__main__":
    main()
